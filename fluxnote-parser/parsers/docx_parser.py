# DOCX file parser

from typing import BinaryIO
from docx import Document


def parse_docx(file: BinaryIO) -> str:
    """
    Parse DOCX file and extract text content.
    
    Args:
        file: Binary file object
        
    Returns:
        Extracted text content
    """
    doc = Document(file)
    
    # Extract text from all paragraphs
    paragraphs = [para.text for para in doc.paragraphs]
    
    # Extract text from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            table_text.append(' | '.join(row_text))
    
    # Combine all text
    all_text = paragraphs + table_text
    return '\n'.join(all_text)
